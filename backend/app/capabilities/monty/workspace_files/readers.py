"""File readers used by the workspace-scoped OS file capability."""

from __future__ import annotations

import csv
import io
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from xml.etree import ElementTree


class FileReadError(ValueError):
    """Raised when a file cannot be converted into useful text."""


@dataclass(slots=True)
class FileReadResult:
    content: str
    file_type: str
    metadata: dict[str, object] = field(default_factory=dict)


_TEXT_EXTENSIONS = {
    "",
    ".bash",
    ".bat",
    ".c",
    ".cfg",
    ".conf",
    ".cpp",
    ".cs",
    ".css",
    ".csv",
    ".dockerfile",
    ".env",
    ".go",
    ".graphql",
    ".h",
    ".html",
    ".ini",
    ".java",
    ".js",
    ".json",
    ".jsonl",
    ".jsx",
    ".log",
    ".md",
    ".mermaid",
    ".mjs",
    ".mmd",
    ".py",
    ".rb",
    ".rs",
    ".scss",
    ".sh",
    ".sql",
    ".svg",
    ".toml",
    ".ts",
    ".tsx",
    ".txt",
    ".xml",
    ".yaml",
    ".yml",
}


def read_file_content(path: Path) -> FileReadResult:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".xlsx":
        return _read_xlsx(path)
    if suffix == ".csv":
        return _read_csv(path)
    if suffix in _TEXT_EXTENSIONS:
        return FileReadResult(_read_text(path), "text")

    try:
        return FileReadResult(_read_text(path), "text")
    except FileReadError as exc:
        raise FileReadError(
            f"Unsupported file type for text extraction: {suffix or 'extensionless'}"
        ) from exc


def _read_text(path: Path) -> str:
    raw = path.read_bytes()
    if b"\x00" in raw[:4096]:
        raise FileReadError("File appears to be binary.")

    encodings = ["utf-8", "utf-8-sig", "cp1252", "latin1", "ascii"]
    for encoding in encodings:
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise FileReadError(f"Could not decode file using: {', '.join(encodings)}")


def _read_csv(path: Path) -> FileReadResult:
    text = _read_text(path)
    sample = text[:8192]
    try:
        dialect = csv.Sniffer().sniff(sample) if sample.strip() else csv.excel
    except csv.Error:
        dialect = csv.excel
    rows = list(csv.reader(io.StringIO(text), dialect=dialect))
    preview = rows[:100]
    rendered = "\n".join(",".join(_csv_cell(cell) for cell in row) for row in preview)
    if len(rows) > len(preview):
        rendered += f"\n... {len(rows) - len(preview)} more row(s)"
    return FileReadResult(
        rendered,
        "csv",
        {"rows": len(rows), "preview_rows": len(preview)},
    )


def _read_pdf(path: Path) -> FileReadResult:
    try:
        import fitz
    except ImportError as exc:
        raise FileReadError("PDF reading requires the pymupdf package.") from exc

    pages: list[str] = []
    with fitz.open(path) as document:
        for page_index, page in enumerate(document, start=1):
            text = page.get_text("text").strip()
            pages.append(f"--- Page {page_index} ---\n{text}")
        page_count = document.page_count
    return FileReadResult("\n\n".join(pages), "pdf", {"pages": page_count})


def _read_docx(path: Path) -> FileReadResult:
    try:
        from docx import Document
    except ImportError:
        return _read_docx_without_dependency(path)

    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    return FileReadResult(
        "\n".join(blocks),
        "docx",
        {"paragraphs": len(document.paragraphs), "tables": len(document.tables)},
    )


def _read_docx_without_dependency(path: Path) -> FileReadResult:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
    root = ElementTree.fromstring(xml)
    paragraphs = []
    for paragraph in root.findall(".//w:p", namespace):
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", namespace))
        if text.strip():
            paragraphs.append(text)
    return FileReadResult("\n".join(paragraphs), "docx", {"paragraphs": len(paragraphs)})


def _read_xlsx(path: Path) -> FileReadResult:
    try:
        import openpyxl
    except ImportError as exc:
        raise FileReadError("XLSX reading requires the openpyxl package.") from exc

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sections: list[str] = []
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        rows = []
        for index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if index > 50:
                rows.append(["..."])
                break
            rows.append(["" if cell is None else str(cell) for cell in row])
        sections.append(f"--- Sheet: {sheet_name} ---\n" + _render_markdown_table(rows))
    return FileReadResult(
        "\n\n".join(sections),
        "xlsx",
        {"sheets": workbook.sheetnames},
    )


def _render_markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return "(empty sheet)"
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * width
    body = normalized[1:]
    table_rows = [header, separator, *body]
    return "\n".join("| " + " | ".join(row) + " |" for row in table_rows)


def _csv_cell(cell: str) -> str:
    if any(character in cell for character in [",", '"', "\n", "\r"]):
        return '"' + cell.replace('"', '""') + '"'
    return cell
